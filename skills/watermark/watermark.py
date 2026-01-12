#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Watermark Tool - 给 PDF、Word、Excel 文件添加水印
"""

import os
import sys
import argparse
from io import BytesIO

# PDF 处理
try:
    from PyPDF2 import PdfReader, PdfWriter
    from reportlab.pdfgen import canvas
    from reportlab.pdfbase import pdfmetrics
    from reportlab.pdfbase.ttfonts import TTFont
    PDF_SUPPORT = True
except ImportError:
    PDF_SUPPORT = False

# Word 处理
try:
    from docx import Document
    from docx.shared import Pt, RGBColor
    DOCX_SUPPORT = True
except ImportError:
    DOCX_SUPPORT = False

# Excel 处理
try:
    from openpyxl import load_workbook
    XLSX_SUPPORT = True
except ImportError:
    XLSX_SUPPORT = False


def register_chinese_font():
    """注册中文字体"""
    if not PDF_SUPPORT:
        return False

    font_paths = [
        "/System/Library/Fonts/PingFang.ttc",
        "/System/Library/Fonts/STHeiti Light.ttc",
        "/Library/Fonts/Arial Unicode.ttf",
        "/usr/share/fonts/truetype/droid/DroidSansFallbackFull.ttf",
        "C:\\Windows\\Fonts\\msyh.ttc",
    ]

    for fp in font_paths:
        if os.path.exists(fp):
            try:
                pdfmetrics.registerFont(TTFont('ChineseFont', fp))
                return True
            except:
                continue
    return False


FONT_REGISTERED = register_chinese_font() if PDF_SUPPORT else False


def create_watermark_pdf(text, page_width, page_height):
    """创建水印 PDF"""
    packet = BytesIO()
    c = canvas.Canvas(packet, pagesize=(page_width, page_height))

    c.setFillColorRGB(0.7, 0.7, 0.7, alpha=0.3)

    if FONT_REGISTERED:
        c.setFont('ChineseFont', 30)
    else:
        c.setFont('Helvetica', 30)

    # 中心水印
    c.saveState()
    c.translate(page_width/2, page_height/2)
    c.rotate(45)
    c.drawCentredString(0, 0, text)
    c.restoreState()

    # 左上角
    c.saveState()
    c.translate(page_width*0.25, page_height*0.75)
    c.rotate(45)
    c.drawCentredString(0, 0, text)
    c.restoreState()

    # 右下角
    c.saveState()
    c.translate(page_width*0.75, page_height*0.25)
    c.rotate(45)
    c.drawCentredString(0, 0, text)
    c.restoreState()

    c.save()
    packet.seek(0)
    return PdfReader(packet)


def add_watermark_pdf(input_path, output_path, text):
    """给 PDF 添加水印"""
    if not PDF_SUPPORT:
        print("错误: 需要安装 PyPDF2 和 reportlab")
        print("运行: pip install PyPDF2 reportlab")
        return False

    try:
        reader = PdfReader(input_path)
        writer = PdfWriter()

        for page in reader.pages:
            page_width = float(page.mediabox.width)
            page_height = float(page.mediabox.height)
            watermark = create_watermark_pdf(text, page_width, page_height)
            page.merge_page(watermark.pages[0])
            writer.add_page(page)

        with open(output_path, 'wb') as f:
            writer.write(f)
        return True
    except Exception as e:
        print(f"PDF 处理失败: {e}")
        return False


def add_watermark_docx(input_path, output_path, text):
    """给 Word 添加水印"""
    if not DOCX_SUPPORT:
        print("错误: 需要安装 python-docx")
        print("运行: pip install python-docx")
        return False

    try:
        doc = Document(input_path)

        for section in doc.sections:
            header = section.header
            p = header.paragraphs[0] if header.paragraphs else header.add_paragraph()
            run = p.add_run(text)
            run.font.size = Pt(36)
            run.font.color.rgb = RGBColor(200, 200, 200)
            p.alignment = 1

        doc.save(output_path)
        return True
    except Exception as e:
        print(f"Word 处理失败: {e}")
        return False


def add_watermark_xlsx(input_path, output_path, text):
    """给 Excel 添加水印"""
    if not XLSX_SUPPORT:
        print("错误: 需要安装 openpyxl")
        print("运行: pip install openpyxl")
        return False

    try:
        wb = load_workbook(input_path)

        for sheet in wb.worksheets:
            sheet.oddHeader.center.text = text
            sheet.oddHeader.center.size = 24
            sheet.oddHeader.center.color = "C0C0C0"
            sheet.evenHeader.center.text = text
            sheet.evenHeader.center.size = 24
            sheet.evenHeader.center.color = "C0C0C0"

        wb.save(output_path)
        return True
    except Exception as e:
        print(f"Excel 处理失败: {e}")
        return False


def process_file(file_path, text, output_dir=None, overwrite=False):
    """处理单个文件"""
    ext = os.path.splitext(file_path)[1].lower()

    if output_dir:
        os.makedirs(output_dir, exist_ok=True)
        output_path = os.path.join(output_dir, os.path.basename(file_path))
    elif overwrite:
        output_path = file_path
        # 先备份
        temp_path = file_path + ".tmp"
        os.rename(file_path, temp_path)
        file_path = temp_path
    else:
        base, extension = os.path.splitext(file_path)
        output_path = f"{base}_watermarked{extension}"

    success = False

    if ext == '.pdf':
        success = add_watermark_pdf(file_path, output_path, text)
    elif ext == '.docx':
        success = add_watermark_docx(file_path, output_path, text)
    elif ext == '.xlsx':
        success = add_watermark_xlsx(file_path, output_path, text)
    else:
        print(f"不支持的文件类型: {ext}")
        return False

    # 清理临时文件
    if overwrite and not output_dir:
        if success:
            os.remove(file_path)
        else:
            os.rename(file_path, output_path.replace(".tmp", ""))

    return success


def process_directory(dir_path, text, output_dir=None, overwrite=False):
    """处理目录中的所有文件"""
    supported_exts = {'.pdf', '.docx', '.xlsx'}
    results = {'success': 0, 'failed': 0}

    for root, dirs, files in os.walk(dir_path):
        # 跳过隐藏目录
        dirs[:] = [d for d in dirs if not d.startswith('.')]

        for f in files:
            if f.startswith('~$') or f.startswith('.'):
                continue

            ext = os.path.splitext(f)[1].lower()
            if ext not in supported_exts:
                continue

            file_path = os.path.join(root, f)

            # 计算输出目录（保持目录结构）
            if output_dir:
                rel_path = os.path.relpath(root, dir_path)
                file_output_dir = os.path.join(output_dir, rel_path)
            else:
                file_output_dir = None

            print(f"处理: {f}")
            if process_file(file_path, text, file_output_dir, overwrite):
                print(f"  ✓ 完成")
                results['success'] += 1
            else:
                print(f"  ✗ 失败")
                results['failed'] += 1

    return results


def main():
    parser = argparse.ArgumentParser(
        description='给 PDF、Word、Excel 文件添加水印',
        formatter_class=argparse.RawDescriptionHelpFormatter,
        epilog='''
示例:
  %(prog)s -t "机密文件" file.pdf
  %(prog)s -t "内部使用" -d ./documents
  %(prog)s -t "仅供参考" -d ./docs -o ./watermarked
  %(prog)s -t "禁止外传" -d ./docs --overwrite
        '''
    )

    parser.add_argument('path', nargs='?', help='文件或目录路径')
    parser.add_argument('-t', '--text', required=True, help='水印文字')
    parser.add_argument('-d', '--directory', help='处理整个目录')
    parser.add_argument('-o', '--output', help='输出目录')
    parser.add_argument('--overwrite', action='store_true', help='覆盖原文件')

    args = parser.parse_args()

    if args.directory:
        target = args.directory
    elif args.path:
        target = args.path
    else:
        parser.print_help()
        sys.exit(1)

    if not os.path.exists(target):
        print(f"错误: 路径不存在 - {target}")
        sys.exit(1)

    if os.path.isfile(target):
        if process_file(target, args.text, args.output, args.overwrite):
            print("✓ 水印添加完成")
        else:
            print("✗ 水印添加失败")
            sys.exit(1)
    else:
        results = process_directory(target, args.text, args.output, args.overwrite)
        print(f"\n完成: {results['success']} 成功, {results['failed']} 失败")


if __name__ == '__main__':
    main()
